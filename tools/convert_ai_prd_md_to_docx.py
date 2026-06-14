import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path("/Users/jelly/Documents/New project")
SRC = ROOT / "AI社区平台竞品调研分析与产品需求PRD_V1.0.md"
OUT = ROOT / "AI社区平台竞品调研分析与产品需求PRD_V1.0_飞书上传版.docx"


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\((https?://[^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"\[([^\]]+)\]\(<([^>]+)>\)", r"\1（\2）", text)
    return text


def split_table_row(line: str):
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_inline(cell.strip()) for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def add_markdown_table(doc: Document, rows):
    if not rows:
        return
    parsed = [split_table_row(r) for r in rows if not is_table_separator(r)]
    if not parsed:
        return
    max_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = "Table Grid"
    for row_idx, row in enumerate(parsed):
        for col_idx in range(max_cols):
            text = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, code: str):
    if not code.strip():
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code.rstrip("\n"))
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    paragraph.paragraph_format.space_after = Pt(6)


def add_image(doc: Document, alt: str, raw_path: str):
    img_path = raw_path.strip()
    if img_path.startswith("<") and img_path.endswith(">"):
        img_path = img_path[1:-1]
    if not os.path.isabs(img_path):
        img_path = str((ROOT / img_path).resolve())
    paragraph = doc.add_paragraph()
    if alt:
        run = paragraph.add_run(alt)
        run.bold = True
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[图片缺失] {img_path}")
        return
    try:
        doc.add_picture(img_path, width=Inches(6.3))
    except Exception as exc:
        doc.add_paragraph(f"[图片插入失败] {img_path}: {exc}")
    doc.add_paragraph()


def apply_base_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    for i in range(1, 6):
        style_name = f"Heading {i}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style.font.bold = True


def convert():
    md = SRC.read_text(encoding="utf-8")
    lines = md.splitlines()
    doc = Document()
    apply_base_styles(doc)

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.strip() == "---":
            paragraph = doc.add_paragraph()
            paragraph.add_run("—" * 28)
            i += 1
            continue

        image_match = re.match(r"!\[([^\]]*)\]\((.+)\)\s*$", line.strip())
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if line.lstrip().startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_lines = [line, lines[i + 1]]
            i += 2
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 5)
            doc.add_heading(clean_inline(heading_match.group(2)), level=level)
            i += 1
            continue

        if line.startswith(">"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(clean_inline(line.lstrip("> ").strip()))
            run.italic = True
            i += 1
            continue

        bullet_match = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet_match:
            doc.add_paragraph(clean_inline(bullet_match.group(1)), style="List Bullet")
            i += 1
            continue

        num_match = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if num_match:
            doc.add_paragraph(clean_inline(num_match.group(2)), style="List Number")
            i += 1
            continue

        doc.add_paragraph(clean_inline(line.strip()))
        i += 1

    if in_code and code_lines:
        add_code_block(doc, "\n".join(code_lines))

    doc.save(OUT)


if __name__ == "__main__":
    convert()
    print(OUT)
